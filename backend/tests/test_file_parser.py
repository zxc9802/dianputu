"""Tests for the file_parser service."""

from __future__ import annotations

import io
from app.services.file_parser import ParseResult, parse_document, MAX_EXTRACTED_TEXT_LENGTH


# ---------------------------------------------------------------------------
# Helper: create real .docx bytes in memory
# ---------------------------------------------------------------------------

def _make_docx_bytes(paragraphs: list[str], table_data: list[list[str]] | None = None) -> bytes:
    """Create a minimal .docx file in memory with the given paragraphs and optional table."""
    from docx import Document

    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    if table_data:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for row_index, row_data in enumerate(table_data):
            for col_index, cell_text in enumerate(row_data):
                table.rows[row_index].cells[col_index].text = cell_text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Helper: create real .xlsx bytes in memory
# ---------------------------------------------------------------------------

def _make_xlsx_bytes(sheets: dict[str, list[list[str | int | None]]]) -> bytes:
    """Create a minimal .xlsx file in memory with named sheets."""
    from openpyxl import Workbook

    wb = Workbook()
    first = True
    for sheet_name, rows in sheets.items():
        if first:
            ws = wb.active
            ws.title = sheet_name
            first = False
        else:
            ws = wb.create_sheet(sheet_name)
        for row in rows:
            ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Helper: create real PDF bytes in memory
# ---------------------------------------------------------------------------

def _make_pdf_bytes(pages: list[str]) -> bytes:
    """Create a minimal PDF file in memory with text pages (no external deps)."""
    # Build a minimal valid PDF with raw syntax.
    objects: list[bytes] = []
    # Object 1: Catalog
    objects.append(b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n")
    # Object 2: Pages (placeholder, will be replaced)
    page_obj_ids: list[int] = []
    next_id = 3
    # Reserve id 2 for Pages object; build page objects first.
    page_objects: list[bytes] = []
    for page_text in pages:
        page_id = next_id
        content_id = next_id + 1
        font_id = next_id + 2
        next_id += 3
        page_obj_ids.append(page_id)
        # Font object
        font_obj = f"{font_id} 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n".encode()
        # Content stream
        escaped = page_text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode()
        content_obj = f"{content_id} 0 obj\n<< /Length {len(stream)} >>\nstream\n".encode() + stream + b"\nendstream\nendobj\n"
        # Page object
        page_obj = f"{page_id} 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents {content_id} 0 R /Resources << /Font << /F1 {font_id} 0 R >> >> >>\nendobj\n".encode()
        page_objects.extend([page_obj, content_obj, font_obj])
    # Build Pages object
    kids = " ".join(f"{pid} 0 R" for pid in page_obj_ids)
    pages_obj = f"2 0 obj\n<< /Type /Pages /Kids [{kids}] /Count {len(page_obj_ids)} >>\nendobj\n".encode()
    objects.append(pages_obj)
    objects.extend(page_objects)
    # Assemble PDF
    body = b"%PDF-1.4\n"
    offsets: list[int] = []
    for obj in objects:
        offsets.append(len(body))
        body += obj
    xref_offset = len(body)
    xref = f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode()
    for off in offsets:
        xref += f"{off:010d} 00000 n \n".encode()
    xref += f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode()
    return body + xref


# ===========================================================================
# Word (.docx) tests
# ===========================================================================

class TestParseDocx:
    def test_extract_paragraphs(self):
        data = _make_docx_bytes(["积雪草修护精华", "核心成分：积雪草提取物", "功效：舒缓修护"])
        result = parse_document(data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "产品说明.docx")
        assert result.parser_used == "docx"
        assert "积雪草修护精华" in result.text
        assert "核心成分" in result.text
        assert "功效" in result.text
        assert not result.warnings

    def test_extract_tables(self):
        data = _make_docx_bytes(
            ["成分表"],
            table_data=[
                ["成分名", "功效"],
                ["透明质酸", "补水保湿"],
                ["烟酰胺", "美白提亮"],
            ],
        )
        result = parse_document(data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "成分表.docx")
        assert result.parser_used == "docx"
        assert "透明质酸" in result.text
        assert "烟酰胺" in result.text
        assert "表格 1" in result.text

    def test_empty_docx(self):
        data = _make_docx_bytes([])
        result = parse_document(data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "空.docx")
        assert result.parser_used == "docx"
        assert result.text == ""
        assert any("未提取到" in w for w in result.warnings)

    def test_extension_fallback(self):
        """When content_type is generic, should use .docx extension to identify parser."""
        data = _make_docx_bytes(["测试内容"])
        result = parse_document(data, "application/octet-stream", "测试.docx")
        assert result.parser_used == "docx"
        assert "测试内容" in result.text


# ===========================================================================
# Excel (.xlsx) tests
# ===========================================================================

class TestParseXlsx:
    def test_single_sheet(self):
        data = _make_xlsx_bytes({"成分表": [["成分", "含量"], ["积雪草", "5%"], ["透明质酸", "2%"]]})
        result = parse_document(data, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "成分.xlsx")
        assert result.parser_used == "xlsx"
        assert "积雪草" in result.text
        assert "透明质酸" in result.text
        assert result.page_count == 1

    def test_multiple_sheets(self):
        data = _make_xlsx_bytes({
            "成分": [["名称", "功效"], ["积雪草", "舒缓"]],
            "检测数据": [["指标", "数值"], ["保湿力", "92%"]],
        })
        result = parse_document(data, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "报告.xlsx")
        assert result.parser_used == "xlsx"
        assert "积雪草" in result.text
        assert "保湿力" in result.text
        assert result.page_count == 2

    def test_empty_xlsx(self):
        data = _make_xlsx_bytes({"Sheet1": []})
        result = parse_document(data, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "空.xlsx")
        assert result.parser_used == "xlsx"
        assert any("未从 Excel" in w for w in result.warnings)


# ===========================================================================
# PDF tests
# ===========================================================================

class TestParsePdf:
    def test_single_page(self):
        data = _make_pdf_bytes(["CICA Repair Serum Product Info"])
        result = parse_document(data, "application/pdf", "产品介绍.pdf")
        assert result.parser_used == "pdf"
        assert "CICA" in result.text
        assert result.page_count == 1

    def test_multi_page(self):
        data = _make_pdf_bytes(["Page 1: Ingredients", "Page 2: Test Results"])
        result = parse_document(data, "application/pdf", "报告.pdf")
        assert result.parser_used == "pdf"
        assert "Ingredients" in result.text
        assert "Test Results" in result.text
        assert result.page_count == 2


# ===========================================================================
# Plain text tests
# ===========================================================================

class TestParseText:
    def test_utf8_text(self):
        data = "产品名称：修护精华\n核心成分：积雪草".encode("utf-8")
        result = parse_document(data, "text/plain", "说明.txt")
        assert result.parser_used == "text"
        assert "修护精华" in result.text

    def test_csv(self):
        data = "成分,含量\n积雪草,5%\n透明质酸,2%".encode("utf-8")
        result = parse_document(data, "text/csv", "成分表.csv")
        assert result.parser_used == "text"
        assert "积雪草" in result.text

    def test_gbk_encoding(self):
        data = "产品名称：修护精华".encode("gbk")
        result = parse_document(data, "text/plain", "说明.txt")
        assert result.parser_used == "text"
        assert "修护精华" in result.text


# ===========================================================================
# Unsupported format tests
# ===========================================================================

class TestUnsupportedFormats:
    def test_doc_legacy(self):
        result = parse_document(b"fake-doc-data", "application/msword", "旧文件.doc")
        assert result.text == ""
        assert any(".doc 格式暂不支持" in w for w in result.warnings)

    def test_xls_legacy(self):
        result = parse_document(b"fake-xls-data", "application/vnd.ms-excel", "旧表格.xls")
        assert result.text == ""
        assert any(".xls 格式暂不支持" in w for w in result.warnings)

    def test_unknown_format(self):
        result = parse_document(b"random-binary", "application/x-unknown", "未知.bin")
        assert result.text == ""
        assert any("无法识别" in w for w in result.warnings)


# ===========================================================================
# Edge case tests
# ===========================================================================

class TestEdgeCases:
    def test_empty_data(self):
        result = parse_document(b"", "application/pdf", "空.pdf")
        assert result.text == ""
        assert any("内容为空" in w for w in result.warnings)

    def test_corrupted_docx(self):
        result = parse_document(b"not-a-real-docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "损坏.docx")
        assert result.parser_used == "docx"
        assert any("解析失败" in w for w in result.warnings)

    def test_text_truncation(self):
        long_text = "测" * (MAX_EXTRACTED_TEXT_LENGTH + 1000)
        data = long_text.encode("utf-8")
        result = parse_document(data, "text/plain", "超长.txt")
        assert len(result.text) <= MAX_EXTRACTED_TEXT_LENGTH + 50  # allow for truncation marker
        assert "已截断" in result.text
