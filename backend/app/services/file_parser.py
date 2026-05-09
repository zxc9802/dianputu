"""Document parsing service.

Extracts plain text from uploaded documents (Word, Excel, PDF, plain text)
so the content can be sent to the LLM for structured product information extraction.

Supported formats:
- .docx  (python-docx)
- .xlsx  (openpyxl)
- .pdf   (PyMuPDF / fitz)
- .txt / .csv  (direct UTF-8 decode)
- .doc / .xls  (unsupported, returns warning)
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path

logger = logging.getLogger("file_parser")

# Maximum characters to keep after extraction to avoid exceeding LLM token limits.
MAX_EXTRACTED_TEXT_LENGTH = 20_000

# MIME type mappings for binary document formats.
DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
XLSX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}
DOC_MIME_TYPES = {
    "application/msword",
}
XLS_MIME_TYPES = {
    "application/vnd.ms-excel",
}
PDF_MIME_TYPES = {
    "application/pdf",
}
TEXT_MIME_PREFIXES = ("text/",)

# File extension fallback when MIME type is generic (application/octet-stream).
EXTENSION_TO_PARSER: dict[str, str] = {
    ".docx": "docx",
    ".doc": "doc_legacy",
    ".xlsx": "xlsx",
    ".xls": "xls_legacy",
    ".pdf": "pdf",
    ".txt": "text",
    ".csv": "text",
}


@dataclass
class ParseResult:
    """Result of parsing a document."""

    text: str = ""
    parser_used: str = "unsupported"
    page_count: int = 0
    warnings: list[str] = field(default_factory=list)


def parse_document(data: bytes, content_type: str, filename: str) -> ParseResult:
    """Parse a document and return extracted plain text.

    This is the main entry point. It dispatches to the appropriate parser
    based on content_type and filename extension.
    """
    if not data:
        return ParseResult(warnings=["文件内容为空"])

    parser_type = _identify_parser(content_type, filename)

    try:
        if parser_type == "docx":
            return _parse_docx(data, filename)
        if parser_type == "xlsx":
            return _parse_xlsx(data, filename)
        if parser_type == "pdf":
            return _parse_pdf(data, filename)
        if parser_type == "text":
            return _parse_text(data, filename)
        if parser_type == "doc_legacy":
            return ParseResult(
                warnings=[f"{filename}: .doc 格式暂不支持，请转换为 .docx 后重新上传。"],
            )
        if parser_type == "xls_legacy":
            return ParseResult(
                warnings=[f"{filename}: .xls 格式暂不支持，请转换为 .xlsx 后重新上传。"],
            )
    except Exception as exc:
        logger.warning("file parse failed file=%s parser=%s error=%s", filename, parser_type, exc)
        return ParseResult(
            parser_used=parser_type,
            warnings=[f"{filename}: 解析失败 ({exc})"],
        )

    return ParseResult(warnings=[f"{filename}: 无法识别的文件格式 ({content_type})"])


def _identify_parser(content_type: str, filename: str) -> str:
    """Determine which parser to use based on MIME type and file extension."""
    ct = (content_type or "").strip().lower().split(";")[0].strip()

    if ct in DOCX_MIME_TYPES:
        return "docx"
    if ct in XLSX_MIME_TYPES:
        return "xlsx"
    if ct in PDF_MIME_TYPES:
        return "pdf"
    if ct in DOC_MIME_TYPES:
        return "doc_legacy"
    if ct in XLS_MIME_TYPES:
        return "xls_legacy"
    if any(ct.startswith(prefix) for prefix in TEXT_MIME_PREFIXES):
        return "text"

    # Fallback: use file extension when MIME type is generic.
    ext = Path(filename).suffix.lower()
    return EXTENSION_TO_PARSER.get(ext, "unsupported")


def _truncate(text: str) -> str:
    """Truncate text to MAX_EXTRACTED_TEXT_LENGTH."""
    if len(text) <= MAX_EXTRACTED_TEXT_LENGTH:
        return text
    return text[:MAX_EXTRACTED_TEXT_LENGTH] + "\n\n…（内容过长，已截断）"


# ---------------------------------------------------------------------------
# Word (.docx)
# ---------------------------------------------------------------------------

def _parse_docx(data: bytes, filename: str) -> ParseResult:
    """Extract text from a .docx file using python-docx."""
    from docx import Document

    doc = Document(BytesIO(data))
    parts: list[str] = []
    warnings: list[str] = []

    # Extract paragraphs.
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract tables.
    for table_index, table in enumerate(doc.tables):
        table_rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cells)
            if row_text.strip(" |"):
                table_rows.append(row_text)
        if table_rows:
            parts.append(f"\n[表格 {table_index + 1}]")
            parts.extend(table_rows)

    if not parts:
        warnings.append(f"{filename}: 文档中未提取到有效文本内容")

    return ParseResult(
        text=_truncate("\n".join(parts)),
        parser_used="docx",
        page_count=1,  # python-docx doesn't expose page count
        warnings=warnings,
    )


# ---------------------------------------------------------------------------
# Excel (.xlsx)
# ---------------------------------------------------------------------------

def _parse_xlsx(data: bytes, filename: str) -> ParseResult:
    """Extract text from a .xlsx file using openpyxl."""
    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    warnings: list[str] = []
    sheet_count = 0

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        sheet_count += 1
        sheet_rows: list[str] = []

        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell).strip() if cell is not None else "" for cell in row]
            row_text = " | ".join(cells)
            if row_text.strip(" |"):
                sheet_rows.append(row_text)

        if sheet_rows:
            parts.append(f"\n[工作表: {sheet_name}]")
            parts.extend(sheet_rows)

    wb.close()

    if not parts:
        warnings.append(f"{filename}: 未从 Excel 中提取到有效内容")

    return ParseResult(
        text=_truncate("\n".join(parts)),
        parser_used="xlsx",
        page_count=sheet_count,
        warnings=warnings,
    )


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _parse_pdf(data: bytes, filename: str) -> ParseResult:
    """Extract text from a PDF file using PyMuPDF (fitz)."""
    import fitz  # PyMuPDF

    doc = fitz.open(stream=data, filetype="pdf")
    parts: list[str] = []
    warnings: list[str] = []
    page_count = doc.page_count

    for page_index in range(page_count):
        page = doc.load_page(page_index)
        page_text = page.get_text("text").strip()
        if page_text:
            parts.append(page_text)

    doc.close()

    if not parts and page_count > 0:
        warnings.append(
            f"{filename}: PDF 共 {page_count} 页但未提取到文本（可能是扫描件/图片型 PDF，建议用产品图上传并让 AI 识图）"
        )

    return ParseResult(
        text=_truncate("\n\n".join(parts)),
        parser_used="pdf",
        page_count=page_count,
        warnings=warnings,
    )


# ---------------------------------------------------------------------------
# Plain text (.txt, .csv)
# ---------------------------------------------------------------------------

def _parse_text(data: bytes, filename: str) -> ParseResult:
    """Decode plain text files."""
    warnings: list[str] = []

    # Try UTF-8 first, then GBK (common for Chinese documents), then latin-1 as fallback.
    text = ""
    for encoding in ("utf-8", "gbk", "latin-1"):
        try:
            text = data.decode(encoding)
            break
        except (UnicodeDecodeError, ValueError):
            continue

    if not text:
        warnings.append(f"{filename}: 无法解码文件内容")

    return ParseResult(
        text=_truncate(text.strip()),
        parser_used="text",
        page_count=1,
        warnings=warnings,
    )
